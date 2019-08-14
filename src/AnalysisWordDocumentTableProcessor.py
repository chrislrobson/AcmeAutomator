####################################################################################################################
# Python Qt5 Testbed Tester Analysis Word Document Table Processor
# MODULE:  AnalysisWordDocumentTableProcessor
# Author: Christopher Robson
# Copyright by:  Christopher Robson
# Copyright date: 01Jan2016
# This software is not FREE!  Use or destribution of the software system and its
# subsystem modules, libraries, configuration file and "seed" file without the
# express permission of the author is strictly PROHIBITED!
# FUNCTION:  Module analysis processing of seeded data
####################################################################################################################
from docx import Document
from docx.shared import Inches
import os, sys, stat
#------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------
# Home grown
#------------------------------------------------------------------------------------------------------------------
from Globals import *
from Exceptions import *
from SeedDictionary import SeedCommandDictionaryProcessor
from SeedCommandlinePreprocessor import SeedCommandlinePreprocessor
#------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------
# Report Word Document Processor
#-----------------------------------------------------------------------
class WordDocumentTableProcessor:
  "Word Document Processor"
  def __init__( self, parent = None ):
    self.name = "Word Document Processor"
    self.parent = parent
    self.cmd_list = parent.cmd_list
    self.report_file_list = parent.report_file_list
  def process_word_document_table(self, data ):
    document = Document( )
    document.add_heading( 'Test Plan Report', 1 )
    document.add_heading( 'FAIL TABLE', level = 2 )
    fail_table = document.add_table( rows = 1, cols = 6 )
    fail_table.style = "TableGrid"
    hdr_cells = fail_table.rows[0].cells
    hdr_cells[0].text = 'Testcase'
    hdr_cells[1].text = 'Pass/Fail'
    hdr_cells[2].text = 'Error File or Failed Command'
    hdr_cells[3].text = 'Line Numbers'
    hdr_cells[4].text = 'Data Expected'
    hdr_cells[5].text = 'Report File'
    for item in data.prompt_keys:
      if "Failed" in item[2]:
        row_cells = fail_table.add_row( ).cells
        row_cells[0].text = item[1]
        row_cells[1].text = item[2]
        row_cells[2].text = item[3]
        row_cells[3].text = "{} thru {}".format( item[5], item[6] )
        row_cells[4].text = item[0]
        row_cells[5].text = item[4]
    document.add_page_break( )
    document.add_heading( 'PASS TABLE', level = 2 )
    pass_table = document.add_table( rows = 1, cols = 6 )
    pass_table.style = "TableGrid"
    hdr_cells = pass_table.rows[0].cells
    hdr_cells[0].text = 'Testcase'
    hdr_cells[1].text = 'Pass/Fail'
    hdr_cells[2].text = 'Command'
    hdr_cells[3].text = 'Line Numbers'
    hdr_cells[4].text = 'Data Expected'
    hdr_cells[5].text = 'Report File'
    for item in data.prompt_keys:
      if "Passed" in item[2]:
        row_cells = pass_table.add_row( ).cells
        row_cells[0].text = item[1]
        row_cells[1].text = item[2]
        row_cells[2].text = item[3]
        row_cells[3].text = "{} thru {}".format( item[5], item[6] )
        row_cells[4].text = item[0]
        row_cells[5].text = item[4]
    document.save( self.cmd_list[SeedCommandDictionaryProcessor.archivefilename] + ".docx" )
    self.filename_chmod = self.cmd_list[SeedCommandDictionaryProcessor.archivefilename] + ".docx"
    os.chmod( self.filename_chmod, stat.S_IRWXU | stat.S_IRWXG | stat.S_IRWXO )
    self.report_file_list.append( self.cmd_list[SeedCommandDictionaryProcessor.archivefilename] + ".docx" )
    return()
#####################################################################################################################
