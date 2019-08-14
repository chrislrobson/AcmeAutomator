####################################################################################################################
# Python Qt5 Testbed Tester Analysis Word Document Processor
# MODULE:  AnalysisWordDocumentProcessor
# Author: Christopher Robson
# Copyright by:  Christopher Robson
# Copyright date: 01Jan2016
# This software is not FREE!  Use or destribution of the software system and its
# subsystem modules, libraries, configuration file and "seed" file without the
# express permission of the author is strictly PROHIBITED!
# FUNCTION:  Module analysis processing of seeded data
####################################################################################################################
from docx import Document
from docx.shared import Inches
import os
#------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------
# Home grown
#------------------------------------------------------------------------------------------------------------------
from Globals import *
from Exceptions import *
from SeedDictionary import SeedCommandDictionaryProcessor
from SeedCommandlinePreprocessor import SeedCommandlinePreprocessor
#------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------
# Report Word Document Processor
#------------------------------------------------------------------------------------------------------------------
class WordDocumentProcessor:
  "Word Document Processor"
  def __init__( self, parent = None ):
    self.name = "Word Document Processor"
    self.parent = parent
    self.report_file_list = parent.report_file_list
  #----------------------------------------------------------------------------------------------------------------
  def process_word_document(self, data ):
    pass
  #----------------------------------------------------------------------------------------------------------------
  def automatically_combine_word_documents( self, target_path, concatenated_files ):
    self.concatenated_files = concatenated_files
    self.target_path = target_path
    self.empty_file = Document()
    self.empty_file.add_heading( 'Test Plan Combined Final Report', 1 )
    self.empty_file.save( self.target_path + "/working_document.docx" )
    try:
      self.working_document = Document( self.target_path + "/working_document.docx" )
    except Exception as error:
      print( "BUMMER {}".format( error ) )
    for file in self.report_file_list:
      sub_doc = Document( file )
      #sub_doc.add_page_break()
      try:
        for element in sub_doc._document_part.body._element:
          self.working_document._document_part.body._element.append( element )
      except Exception as error:
        print( "BUMMER {}".format( error ) )
    self.working_document.save( self.concatenated_files )
    os.remove( self.target_path + "/working_document.docx" )
    return()
  #----------------------------------------------------------------------------------------------------------------
  def combine_word_documents( self, target_path, concatenated_files, source_path, seed_file_list ):
    self.concatenated_files = concatenated_files
    self.seed_file_list = seed_file_list
    self.target_path = target_path
    self.source_path = source_path
    self.working_document = Document( self.target_path + "/working_document.docx" )
    for file in self.seed_file_list:
      sub_doc = Document( file )
      sub_doc.add_page_break()
      for element in sub_doc._document_part.body._element:
        self.working_document._document_part.body._element.append( element )
    self.working_document.save( self.concatenated_files )
    os.remove( self.target_path + "/working_document.docx" )
    return()
#####################################################################################################################
